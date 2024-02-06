# Create your views here.

from django.shortcuts import render
from django.http import HttpResponse
from django.shortcuts import render


from docx import Document

#importing the local preprocessing function

from .ngram_similarity import compute_ngram_similarity,generate_n_grams
from .fingerprint_similarity import compute_fingerprint_similarity,generate_fingerprint_sets
from .word_similarity import compute_word_similarity,generate_word_sim_sets
from .tfidf_similarity import compute_tfidfsimilarity
import time




from .models import thesis_docx
from .predict import predict_lab
#from .forms import DocxUploadForm
from .highlight_paragraphs import highlight_paragraph


def homepage(request):
    return render(request, 'plag_check.html')
    
def thesis_upload_page(request):
    return render(request, 'thesis_upload.html')

def thesis_upload_succesfull_page(request):
    #print('agadi')
    #print(request)
    if request.method == 'POST':
        #print('paxadi')
        #print(request.FILES)

        got_thesis_file = request.FILES.get("thesis_file_docx")
        got_author_name = request.POST.get("author_name")
        document = Document(got_thesis_file)
        input_paragraphs = [[generate_n_grams(para.text.strip()),generate_fingerprint_sets(para.text.strip()),generate_word_sim_sets(para.text.strip()),para.text.strip()] for para in document.paragraphs if para.text.strip() and len(para.text.strip().split())>5]  # Filter out empty or whitespace paragraphs
        #input_paragraphs_json = json.loads(str(input_paragraphs))
        #print(got_thesis_file)
        print(input_paragraphs) 

        #import pprint
        #input_paragraphs_string = pprint.pformat(input_paragraphs)
        #import ast
        #input_paragraphs_after = ast.literal_eval(input_paragraphs_string)
        #if input_paragraphs == input_paragraphs_after:
        #    print(True)
        thesis_data = thesis_docx.objects.create(thesis = got_thesis_file,author = got_author_name,paragraphs = input_paragraphs)
        thesis_data.save()
        return HttpResponse("Thesis Uploaded Sucessfully")
    


def count_total_words(paragraphs):
    total_words = 0
    for paragraph in paragraphs:
        total_words += len(paragraph[3].split())
    return total_words


def count_total_words_one_paragraph(paragraph):
    total_words = 0
    total_words += len(paragraph.split())
    return total_words



def plagiarism_check(request):
    if request.method == 'POST':
        start_time = time.time()
        #form = DocxUploadForm(request.POST, request.FILES)
        #if form.is_valid():
        docx_file = request.FILES.get("thesis_file_docx")
        suspicious_author_name = request.POST.get("author_name")
        #docx_file = request.FILES['docx_file']
        suspicious_docx_file_title_name = docx_file.name
        #print(suspicious_docx_file_title_name)
        #print(suspicious_author_name)
        #print("yesko mathi")
        #docx_file_author_name = 'Inputed_Thesis_Author_Name'
        #print(type(docx_file))
        sus_document = Document(docx_file)

        #print(type(sus_document))
        input_paragraphs = [[generate_n_grams(para.text.strip()),generate_fingerprint_sets(para.text.strip()),generate_word_sim_sets(para.text.strip()),para.text.strip()] for para in sus_document.paragraphs if para.text.strip() and len(para.text.strip().split())>5]  # Filter out empty or whitespace paragraphs
        # paragraphs = get_paragraphs_from_word_file(docx_file)
        #print(input_paragraphs)
        # print(type(paragraphs))
        #return render(request, 'result.html', {'paragraphs': paragraphs})
        

        docx_files = thesis_docx.objects.all()
        #file_names = [docx_file.thesis.name for docx_file in docx_files]
        #print(file_names)
        #print(file_names[0])
        paragraphs_list = []

        for docx_file in docx_files:
            print(docx_file.paragraphs)
            #document = Document(docx_file.thesis)
            #paragraphs = [para.text for para in document.paragraphs]
            #para ={"paragraphs":[[generate_n_grams(para.text.strip()),generate_fingerprint_sets(para.text.strip()),generate_word_sim_sets(para.text.strip()),para.text.strip()] for para in document.paragraphs if para.text.strip() and len(para.text.strip().split())>5],"source":{'name':docx_file.thesis.name,'author':docx_file.author}}  # Filter out empty or whitespace paragraphs
            para ={"paragraphs":docx_file.paragraphs,"source":{'name':docx_file.thesis.name,'author':docx_file.author}}  # Filter out empty or whitespace paragraphs

            #paragraphs = get_paragraphs_from_word_file(document)
            paragraphs_list.append(para)
        
        


        plagiarised_paragraphs =[]
        # Loop through both lists simultaneously and call the function
        for para1 in input_paragraphs:
           
            #paragraphs_list_counter =0
            for para2_group in paragraphs_list:
                max_avg_feature_sim =0
                max_sim_database_paragraph = ''
                is_label_one = 0
                for para2 in para2_group['paragraphs']:
                    #asma two paragraph aaucha
                    #print('--------------------------------------------------------------')
                    #print(para2[0])
                    #print("----------------------------------------------------")
                    ngram_sim = compute_ngram_similarity(para1[0],para2[0])
                    #print(para1[3])
                    #print(para2[3])
                    #print(para1[1])
                    #print(para2[1])
                    finger_sim = compute_fingerprint_similarity(para1[1],para2[1])
                    
                    word_sim = compute_word_similarity(para1[2],para2[2])
                    tfidf_sim = compute_tfidfsimilarity(para1[3],para2[3]) #4th item in list is paragraph 
                    label = predict_lab(ngram_sim,finger_sim,word_sim,tfidf_sim)
                    avg_feature_sim = (ngram_sim + tfidf_sim + finger_sim + word_sim)/4
                    if(avg_feature_sim > max_avg_feature_sim and label == 1):
                        max_avg_feature_sim = avg_feature_sim
                        max_sim_database_paragraph = para2[3]
                    
                    if(label == 1):
                        is_label_one = 1
                if is_label_one ==1 :
                    my_dict = {"paragraph": para1[3] ,"database_paragraph":max_sim_database_paragraph, "source":para2_group['source'],"average_feature_score": max_avg_feature_sim}
                    plagiarised_paragraphs.append(my_dict)
                    #asma lekhne sentence wala code
                    #
                is_label_one = 0
                #paragraphs_list_counter = paragraphs_list_counter + 1
            #         print(para1)
            #         print("-----")
            #         print(para2)
            #         print("label: ")
            #         print(label)
            #     print("changeeeeeeeeeeee1111111")
            # print("changeeeeeeeeeeee22222222")
        #print(plagiarised_paragraphs)
        #return render(request, 'try.html', {'paragraphs_list': paragraphs_list})
        #for highlighting docx
        #output_file_path = "highlighted_document.docx"  # Replace with the desired output file path
        #target_paragraphs = plagiarised_paragraphs
        #process_plagiarized_paragraphs(sus_document, output_file_path, target_paragraphs)
        final_plagiarised_paragraphs = {}
        for entry in plagiarised_paragraphs:
            paragraph = entry['paragraph']
            if paragraph in final_plagiarised_paragraphs:
                if entry['average_feature_score'] > final_plagiarised_paragraphs[paragraph]['average_feature_score']:
                    final_plagiarised_paragraphs[paragraph] = entry
            else:
                final_plagiarised_paragraphs[paragraph] = entry

        #for highlighting docx
        final_plagiarised_paragraphs = list(final_plagiarised_paragraphs.values())
        #print(final_plagiarised_paragraphs)
        total_words_input_docx = count_total_words(input_paragraphs)
        #print(total_words_input_docx)
        # Calculate and store the similarity index for each paragraph
        similarity_data = []
        for entry in final_plagiarised_paragraphs:
            sim_index = round((((count_total_words_one_paragraph(entry['paragraph']))*entry['average_feature_score'])/ total_words_input_docx)*100,2)
            similarity_data.append({'input_paragraph':entry['paragraph'],'source': entry['source'],'database_paragraph':entry['database_paragraph'], 'sim_index': sim_index})
        #print(similarity_data)

        #final wala dictionary aba
        grouped_paragraphs = []
        for entry in similarity_data:
            source = entry['source']
            input_paragraph = entry['input_paragraph']
            database_paragraph = entry['database_paragraph']
            sim_index = entry['sim_index']
            
            # Initialize a flag to False
            contains_source = False
            for single_dict in grouped_paragraphs:
                if 'source' in single_dict and single_dict['source'] == source:
                    single_dict['input_paragraphs'].append(input_paragraph)
                    single_dict['database_paragraphs'].append(database_paragraph)
                    single_dict['percentage'] += sim_index
                    contains_source = True
                    break
            if(contains_source == False):
                grouped_paragraphs.append({'input_paragraphs':[entry['input_paragraph']],'source': entry['source'],'database_paragraphs':[entry['database_paragraph']], 'percentage': sim_index})
        #print("-------------------------")
        #print(grouped_paragraphs)
        #print("----------------------------------")

        #highlight wala
        highlight_paragraph(grouped_paragraphs,sus_document,suspicious_docx_file_title_name,suspicious_author_name)
        end_time = time.time()
        execution_time = end_time - start_time
        minutes, seconds = divmod(execution_time, 60)


        print("--------------------------")
        print(f"Execution time: {int(minutes)} minutes and {round(seconds, 2)} seconds")
        print("--------------------------")


        print(input_paragraphs[0])
        
        #context = {'result_list': similarity_data}
        #return render(request, 'try1.html', context)
        return HttpResponse("Plagiarism Checking Sucessfull.")
    else:
        #form = DocxUploadForm()
        return render(request, 'thesis_upload_for_checking.html')

